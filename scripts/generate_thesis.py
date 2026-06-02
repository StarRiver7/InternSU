# -*- coding: utf-8 -*-
import sys, os
sys.path.insert(0, r"C:\Users\Xiang\.codex\plugins\cache\openai-primary-runtime\documents\26.521.10419\skills\documents\scripts")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
OUTPUT_DIR = r"F:\_Code\Ai\InternSU\docs\documents"
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "InternSU_\u6bd5\u4e1a\u8bbe\u8ba1\u62a5\u544a.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()

for s in doc.sections:
    s.page_width  = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(2.54)
    s.left_margin = s.right_margin = Cm(3.18)

BT = "\u5b8b\u4f53"; BE = "Times New Roman"; HT = "\u9ed1\u4f53"; HE = "Arial"
CODE_F = "Consolas"

def font(run, cn, en, sz, bold=False, color=None):
    run.font.size = Pt(sz); run.bold = bold; run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None: rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>'); rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), cn); rf.set(qn("w:ascii"), en); rf.set(qn("w:hAnsi"), en)
    if color: run.font.color.rgb = color

def pp(text, cn=BT, en=BE, sz=12, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
       sb=0, sa=6, fi=0.74, color=None, ls=None):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format; pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    if fi: pf.first_line_indent = Cm(fi)
    if ls: pf.line_spacing = ls
    r = p.add_run(text); font(r, cn, en, sz, bold, color)
    return p

def heading(text, lv):
    sizes  = {0:22,1:16,2:14,3:12}
    sb_map = {0:24,1:18,2:14,3:10}; sa_map = {0:12,1:10,2:8,3:6}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if lv==0 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(sb_map.get(lv,12)); pf.space_after = Pt(sa_map.get(lv,8))
    pf.first_line_indent = Cm(0)
    r = p.add_run(text)
    font(r, HT, HE, sizes.get(lv,12), True if lv<=2 else False)
    pPr = p._element.get_or_add_pPr()
    ol = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{max(0,lv-1)}"/>')
    old = pPr.find(qn("w:outlineLvl"))
    if old is not None: pPr.remove(old)
    pPr.append(ol)
    return p

def code_block(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before=Pt(2); pf.space_after=Pt(2); pf.left_indent=Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>'))
    r = p.add_run(text); font(r, CODE_F, CODE_F, 9, False, RGBColor(0x33,0x33,0x33))
    return p

def make_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style="Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=""
        r = c.paragraphs[0].add_run(h); font(r,HT,HE,9.5,True,RGBColor(0xFF,0xFF,0xFF))
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc = c._element.get_or_add_tcPr()
        tc.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="2E74B5"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=""
            r=c.paragraphs[0].add_run(str(val)); font(r,BT,BE,9.5,False)
            if ri%2==1:
                tc=c._element.get_or_add_tcPr()
                tc.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F6FC"/>'))
    return t

def brk():
    doc.add_page_break()

print("Functions defined. Building document...")
